#!/usr/bin/env python3
"""
Resume Parser - Extract text from PDF and DOCX files
"""
import sys
import json
import pdfplumber
from docx import Document
import re

def clean_text(text):
    """Clean and normalize extracted text"""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special characters that might cause issues
    text = text.strip()
    return text

def parse_pdf(file_path):
    """Extract text from PDF using pdfplumber"""
    try:
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        if not text.strip():
            return {
                "success": False,
                "error": "No text could be extracted from PDF. It may be an image-based PDF."
            }
        
        return {
            "success": True,
            "text": clean_text(text),
            "pages": len(pdf.pages)
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to parse PDF: {str(e)}"
        }

def parse_docx(file_path):
    """Extract text from DOCX"""
    try:
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        if not text.strip():
            return {
                "success": False,
                "error": "No text found in DOCX file"
            }
        
        return {
            "success": True,
            "text": clean_text(text),
            "pages": len(doc.sections)
        }
    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to parse DOCX: {str(e)}"
        }

def main():
    if len(sys.argv) < 2:
        result = {
            "success": False,
            "error": "No file path provided"
        }
        print(json.dumps(result))
        sys.exit(1)
    
    file_path = sys.argv[1]
    
    # Determine file type and parse accordingly
    if file_path.lower().endswith('.pdf'):
        result = parse_pdf(file_path)
    elif file_path.lower().endswith('.docx'):
        result = parse_docx(file_path)
    else:
        result = {
            "success": False,
            "error": "Unsupported file type. Only PDF and DOCX are supported."
        }
    
    print(json.dumps(result))

if __name__ == "__main__":
    main()
